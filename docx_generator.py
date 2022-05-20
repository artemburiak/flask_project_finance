from docx import Document
import os


class CreateDocument:

    def __init__(self, user_id, username, transactions):
        self.d = Document()

        if not os.path.exists(os.getcwd() + '/export_files'):
            os.mkdir(os.getcwd() + '/export_files')

        self.path = os.getcwd() + '/export_files/' + str(user_id) + '.docx'
        self.username = username

        self.incomes = []
        self.outcomes = []

        for tr in transactions:
            if tr[5] == 1:
                self.incomes.append(tr)
            elif tr[5] == 0:
                self.outcomes.append(tr)

        # print(self.incomes, self.outcomes)
        self.create()



    def create(self):
        self.d.add_heading(self.username, 0)
        self.d.add_heading('Incomes', 1)
        t1 = self.d.add_table(rows=len(self.incomes), cols=3)
        t1.style = 'Table Grid'
        for row in range(len(self.incomes)):
            t1.cell(row, 0).text = str(self.incomes[row][3])
            t1.cell(row, 1).text = self.incomes[row][4]
            t1.cell(row, 2).text = self.incomes[row][2]

        self.d.add_heading('Outcomes', 1)
        t2 = self.d.add_table(rows=len(self.outcomes), cols=3)
        t2.style = 'Table Grid'
        for row in range(len(self.outcomes)):
            t2.cell(row, 0).text = str(self.outcomes[row][3])
            t2.cell(row, 1).text = self.outcomes[row][4]
            t2.cell(row, 2).text = self.outcomes[row][2]

    def save(self):
        self.d.save(self.path)
        return self.path



